"""
src/rag/loaders.py
───────────────────
Read documents off disk and return plain text plus metadata.

Supported out of the box (no extra dependency):
    .md .markdown .txt .rst .csv .tsv .json .log .py

Supported when the optional library is installed:
    .pdf   → needs ``pypdf``        (pip install pypdf)
    .docx  → needs ``python-docx``  (pip install python-docx)

A file whose type needs a missing library is skipped with a clear
warning rather than crashing the ingest — so a company can point the
pipeline at a mixed archive and index everything it can today.
"""

from __future__ import annotations

import csv
import io
import logging
from pathlib import Path
from typing import Iterable, List, Optional

log = logging.getLogger(__name__)

TEXT_SUFFIXES = {
    ".md", ".markdown", ".txt", ".rst", ".log", ".py", ".json",
}
TABULAR_SUFFIXES = {".csv", ".tsv"}
PDF_SUFFIXES = {".pdf"}
DOCX_SUFFIXES = {".docx"}

SUPPORTED_SUFFIXES = (
    TEXT_SUFFIXES | TABULAR_SUFFIXES | PDF_SUFFIXES | DOCX_SUFFIXES
)

# Directories never worth indexing.
SKIP_DIRS = {".git", "__pycache__", ".ipynb_checkpoints", "node_modules",
             ".venv", "venv", "rag_store", "mlruns"}


class Document:
    """A loaded document: its text and where it came from."""

    __slots__ = ("text", "source", "meta")

    def __init__(self, text: str, source: str, meta: Optional[dict] = None):
        self.text = text
        self.source = source
        self.meta = meta or {}

    def __repr__(self) -> str:
        return f"Document(source={self.source!r}, chars={len(self.text)})"


# ── Per-format readers ───────────────────────────────────────────

def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_tabular(path: Path) -> str:
    """Flatten a CSV/TSV into 'col: value' lines so rows are searchable."""
    delim = "\t" if path.suffix.lower() == ".tsv" else ","
    raw = path.read_text(encoding="utf-8", errors="replace")
    reader = csv.reader(io.StringIO(raw), delimiter=delim)
    rows = list(reader)
    if not rows:
        return ""
    header = rows[0]
    lines: List[str] = [f"Table: {path.stem} (columns: {', '.join(header)})"]
    for r in rows[1:]:
        pairs = [f"{h}: {v}" for h, v in zip(header, r) if v != ""]
        if pairs:
            lines.append(" | ".join(pairs))
    return "\n".join(lines)


def _read_pdf(path: Path) -> Optional[str]:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except ImportError:
            log.warning("Skipping %s — install 'pypdf' to index PDFs.",
                        path.name)
            return None
    try:
        reader = PdfReader(str(path))
        parts = []
        for i, page in enumerate(reader.pages):
            txt = page.extract_text() or ""
            if txt.strip():
                parts.append(f"[page {i + 1}]\n{txt}")
        return "\n\n".join(parts)
    except Exception as e:
        log.warning("Failed to read PDF %s: %s", path.name, e)
        return None


def _read_docx(path: Path) -> Optional[str]:
    try:
        import docx  # type: ignore
    except ImportError:
        log.warning("Skipping %s — install 'python-docx' to index .docx.",
                    path.name)
        return None
    try:
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        log.warning("Failed to read DOCX %s: %s", path.name, e)
        return None


# ── Public loaders ───────────────────────────────────────────────

def load_file(path, base: Optional[Path] = None) -> Optional[Document]:
    """Load a single file into a Document, or ``None`` if unsupported/empty."""
    path = Path(path)
    if not path.is_file():
        log.warning("Not a file: %s", path)
        return None

    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        text = _read_text(path)
    elif suffix in TABULAR_SUFFIXES:
        text = _read_tabular(path)
    elif suffix in PDF_SUFFIXES:
        text = _read_pdf(path)
    elif suffix in DOCX_SUFFIXES:
        text = _read_docx(path)
    else:
        log.debug("Unsupported file type, skipping: %s", path.name)
        return None

    if not text or not text.strip():
        return None

    try:
        source = str(path.relative_to(base)) if base else str(path)
    except ValueError:
        source = str(path)
    return Document(text=text, source=source,
                    meta={"suffix": suffix, "path": str(path.resolve())})


def load_dir(directory,
             recursive: bool = True,
             suffixes: Optional[Iterable[str]] = None) -> List[Document]:
    """Load every supported file under ``directory``.

    Args:
        directory: Root directory to walk.
        recursive: Recurse into subdirectories (skipping SKIP_DIRS).
        suffixes:  Restrict to these extensions (e.g. {'.md', '.pdf'}).
                   Defaults to all supported types.
    """
    directory = Path(directory)
    if not directory.is_dir():
        log.warning("Not a directory: %s", directory)
        return []

    allow = {s.lower() for s in suffixes} if suffixes else SUPPORTED_SUFFIXES
    walker = directory.rglob("*") if recursive else directory.glob("*")

    docs: List[Document] = []
    for p in sorted(walker):
        if p.is_dir():
            continue
        if any(part in SKIP_DIRS for part in p.parts):
            continue
        if p.suffix.lower() not in allow:
            continue
        doc = load_file(p, base=directory)
        if doc:
            docs.append(doc)
    return docs
